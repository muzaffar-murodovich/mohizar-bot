from __future__ import annotations

import io

import pytest

from mohizarbot.config import Settings
from mohizarbot.multimodal.document import (
    _MAX_EXTRACTED_CHARS,
    _MAX_PDF_PAGES,
    DocumentProcessor,
    _extract_text,
    supports_mime,
)


@pytest.fixture
def doc_settings() -> Settings:
    return Settings()  # type: ignore[call-arg]


@pytest.fixture
def doc_processor(doc_settings: Settings) -> DocumentProcessor:
    return DocumentProcessor(doc_settings)


@pytest.mark.asyncio
async def test_pdf_text_extraction(doc_processor: DocumentProcessor) -> None:
    """PDF text is extracted using pypdf."""
    try:
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf not installed")

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    result = await doc_processor.process(
        pdf_bytes,
        "application/pdf",
        {"file_name": "test.pdf", "file_id": "pdf1"},
    )

    assert result.kind == "text"
    assert "document_content" in result.text
    assert 'filename="test.pdf"' in result.text
    assert 'source_kind="application"' in result.text


@pytest.mark.asyncio
async def test_pdf_wrapped_in_tags(doc_processor: DocumentProcessor) -> None:
    """PDF output is wrapped in <document_content> tags."""
    try:
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf not installed")

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    result = await doc_processor.process(
        pdf_bytes,
        "application/pdf",
        {"file_name": "doc.pdf", "file_id": "pdf2"},
    )

    assert result.text.startswith("<document_content")
    assert result.text.endswith("</document_content>")


@pytest.mark.asyncio
async def test_pdf_too_many_pages_rejected(doc_processor: DocumentProcessor) -> None:
    """PDF with >100 pages is rejected."""
    try:
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf not installed")

    writer = PdfWriter()
    for _ in range(_MAX_PDF_PAGES + 1):
        writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    with pytest.raises(ValueError, match="too long"):
        await doc_processor.process(
            pdf_bytes,
            "application/pdf",
            {"file_name": "huge.pdf", "file_id": "big1"},
        )


@pytest.mark.asyncio
async def test_docx_text_extraction(doc_processor: DocumentProcessor) -> None:
    """DOCX text is extracted using python-docx."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_paragraph("Hello from DOCX")
    doc.add_paragraph("Second paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    result = await doc_processor.process(
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        {"file_name": "test.docx", "file_id": "docx1"},
    )

    assert result.kind == "text"
    assert "Hello from DOCX" in result.text
    assert "Second paragraph" in result.text
    assert 'filename="test.docx"' in result.text


@pytest.mark.asyncio
async def test_txt_non_utf8_decoded_with_chardet(doc_processor: DocumentProcessor) -> None:
    """Text files with non-UTF-8 encoding are decoded via chardet."""
    # Windows-1251 encoded Russian text — use a longer sample for reliable detection
    text_bytes = ("Привет мир! Как дела? Это тестовый документ. " * 20).encode("windows-1251")

    result = await doc_processor.process(
        text_bytes,
        "text/plain",
        {"file_name": "russian.txt", "file_id": "txt1"},
    )

    assert result.kind == "text"
    assert "Привет мир!" in result.text
    assert 'filename="russian.txt"' in result.text


@pytest.mark.asyncio
async def test_csv_processed_as_text(doc_processor: DocumentProcessor) -> None:
    """CSV files are processed through the text pipeline."""
    csv_bytes = b"name,age,city\nAlice,30,NYC\nBob,25,LA\n"

    result = await doc_processor.process(
        csv_bytes,
        "text/csv",
        {"file_name": "data.csv", "file_id": "csv1"},
    )

    assert result.kind == "text"
    assert "Alice" in result.text
    assert "Bob" in result.text


@pytest.mark.asyncio
async def test_document_too_large_rejected(doc_processor: DocumentProcessor) -> None:
    """Documents > 10MB are rejected."""
    oversized = b"\x00" * (10 * 1024 * 1024 + 1)

    with pytest.raises(ValueError, match="too large"):
        await doc_processor.process(
            oversized,
            "text/plain",
            {"file_name": "huge.txt", "file_id": "big2"},
        )


@pytest.mark.asyncio
async def test_text_truncated_to_max_chars(doc_processor: DocumentProcessor) -> None:
    """Text longer than _MAX_EXTRACTED_CHARS is truncated with warning."""
    long_text = b"x" * (_MAX_EXTRACTED_CHARS + 1000)

    result = await doc_processor.process(
        long_text,
        "text/plain",
        {"file_name": "long.txt", "file_id": "long1"},
    )

    assert len(result.text) <= _MAX_EXTRACTED_CHARS + 500  # accounting for wrapper tags
    assert any("truncated" in w.lower() for w in result.warnings)


@pytest.mark.asyncio
async def test_unsupported_mime_raises(doc_processor: DocumentProcessor) -> None:
    """Unsupported MIME types raise ValueError."""
    with pytest.raises(ValueError, match="Unsupported document MIME"):
        await doc_processor.process(
            b"\x00" * 100,
            "application/zip",
            {"file_name": "bad.zip", "file_id": "zip1"},
        )


def test_supports_mime_pdf() -> None:
    """supports_mime returns True for PDF."""
    assert supports_mime("application/pdf") is True


def test_supports_mime_docx() -> None:
    """supports_mime returns True for DOCX."""
    assert (
        supports_mime("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        is True
    )


def test_supports_mime_text() -> None:
    """supports_mime returns True for text/* types."""
    assert supports_mime("text/plain") is True
    assert supports_mime("text/csv") is True
    assert supports_mime("text/html") is True


def test_supports_mime_non_document() -> None:
    """supports_mime returns False for non-document types."""
    assert supports_mime("audio/ogg") is False
    assert supports_mime("image/jpeg") is False


def test_extract_text_fallback_encoding() -> None:
    """_extract_text falls back to UTF-8 with replace on bad data."""
    text, pages, warnings = _extract_text(b"\xff\xfe\x00\x00" + b"\x00" * 100)
    assert isinstance(text, str)
    assert len(text) > 0
