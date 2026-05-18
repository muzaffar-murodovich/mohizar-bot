from __future__ import annotations

import io
import logging
from typing import TYPE_CHECKING

from mohizarbot.multimodal.base import MultimodalProcessor, ProcessedContent

if TYPE_CHECKING:
    from mohizarbot.config import Settings

logger = logging.getLogger(__name__)

_MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
_MAX_PDF_PAGES = 100
_MAX_EXTRACTED_CHARS = 50_000


class DocumentProcessor(MultimodalProcessor):
    """Processes documents: PDF (pypdf), DOCX (python-docx), TXT/CSV.

    Extracted text is wrapped in <document_content> tags.
    """

    def __init__(self, settings: Settings) -> None:
        self._settings = settings

    async def process(
        self,
        file_bytes: bytes,
        mime_type: str,
        ctx: dict[str, object],
    ) -> ProcessedContent:
        """Extract text from a document file.

        Args:
            file_bytes: Raw document bytes.
            mime_type: MIME type of the document.
            ctx: Must contain `file_name` (str).

        Returns:
            ProcessedContent with kind="text" and extracted content wrapped
            in <document_content> tags.
        """
        warnings: list[str] = []

        if len(file_bytes) > _MAX_FILE_SIZE_BYTES:
            mb = len(file_bytes) / (1024 * 1024)
            raise ValueError(f"Document too large: {mb:.1f}MB (max 10MB)")

        file_name = str(ctx.get("file_name", "document.bin"))

        if mime_type == "application/pdf":
            text, pages, w = _extract_pdf(file_bytes)
            warnings.extend(w)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            text, pages, w = _extract_docx(file_bytes)
            warnings.extend(w)
            pages = str(pages)
        elif mime_type.startswith("text/"):
            text, pages, w = _extract_text(file_bytes)
            warnings.extend(w)
            pages = str(pages)
        else:
            raise ValueError(f"Unsupported document MIME type: {mime_type}")

        # Truncate to max chars
        if len(text) > _MAX_EXTRACTED_CHARS:
            text = text[:_MAX_EXTRACTED_CHARS]
            warnings.append(f"Text truncated to {_MAX_EXTRACTED_CHARS} chars (original was larger)")

        wrapped = (
            f'<document_content filename="{file_name}" pages="{pages}" '
            f'source_kind="{mime_type.split("/")[0]}">\n'
            f"{text}\n"
            f"</document_content>"
        )

        metadata = {
            "source_kind": "document",
            "file_name": file_name,
            "mime_type": mime_type,
            "pages": str(pages),
        }

        return ProcessedContent(
            kind="text",
            text=wrapped,
            metadata=metadata,
            warnings=warnings,
        )


def _extract_pdf(data: bytes) -> tuple[str, str, list[str]]:
    """Extract text from PDF using pypdf."""
    warnings: list[str] = []
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError("pypdf is required for PDF processing") from None

    reader = PdfReader(io.BytesIO(data))
    total_pages = len(reader.pages)

    if total_pages > _MAX_PDF_PAGES:
        raise ValueError(f"PDF too long: {total_pages} pages (max {_MAX_PDF_PAGES})")

    texts: list[str] = []
    pages_extracted = 0
    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            warnings.append(f"Failed to extract text from page {i + 1}")
            page_text = ""
        texts.append(page_text)
        pages_extracted += 1
        # Check cumulative length
        if sum(len(t) for t in texts) > _MAX_EXTRACTED_CHARS * 2:
            warnings.append("PDF text extraction truncated due to size")
            break

    return "\n\n".join(texts), str(pages_extracted), warnings


def _extract_docx(data: bytes) -> tuple[str, str, list[str]]:
    """Extract text from DOCX using python-docx."""
    warnings: list[str] = []
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX processing") from None

    doc = Document(io.BytesIO(data))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_texts: list[str] = []
            for cell in row.cells:
                row_texts.append(cell.text)
            paragraphs.append(" | ".join(row_texts))

    # Check comments
    _has_comment_type = False
    _comment_rel_type: object = None
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as _rt  # noqa: N811

        _has_comment_type = True
        _comment_rel_type = _rt
    except ImportError:
        pass

    comment_count = 0
    if _has_comment_type:
        try:
            comments_part = None
            for rel in doc.part.rels.values():
                if rel.reltype == _comment_rel_type.COMMENTS:  # type: ignore[attr-defined]
                    comments_part = rel.target_part
                    break
            if comments_part is not None:
                from lxml import etree  # type: ignore[import-untyped]

                comments_xml = etree.fromstring(comments_part.blob)
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                for comment_elem in comments_xml.findall(".//w:comment", ns):
                    comment_text_parts = []
                    for t_elem in comment_elem.findall(".//w:t", ns):
                        if t_elem.text:
                            comment_text_parts.append(t_elem.text)
                    comment_body = "".join(comment_text_parts)
                    author = comment_elem.get("{" + ns["w"] + "}author", "unknown")
                    paragraphs.append(f"[Comment by {author}]: {comment_body}")
                    comment_count += 1
        except Exception:
            pass

    text = "\n\n".join(paragraphs)
    return text, str(len(doc.paragraphs)), warnings


def _extract_text(data: bytes) -> tuple[str, str, list[str]]:
    """Extract text from TXT/CSV with charset detection."""
    warnings: list[str] = []
    try:
        import chardet
    except ImportError:
        raise RuntimeError("chardet is required for text file processing") from None

    detection = chardet.detect(data)
    detected_encoding = detection.get("encoding") or "utf-8"
    confidence = detection.get("confidence", 0.0)

    # Try detected encoding first, even at lower confidence
    text = ""
    encoding = detected_encoding
    try:
        text = data.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        if confidence < 0.5:
            warnings.append(f"Low confidence ({confidence:.1%}) for encoding {encoding}")
        try:
            text = data.decode("utf-8", errors="replace")
            encoding = "utf-8"
            warnings.append("Falling back to UTF-8 with replacement chars")
        except Exception:
            text = data.decode("latin-1", errors="replace")
            encoding = "latin-1"
            warnings.append("Falling back to latin-1 with replacement chars")
    else:
        if confidence < 0.5:
            warnings.append(f"Low confidence ({confidence:.1%}) for encoding {encoding}")

    lines = text.splitlines()
    pages = str(max(1, len(lines) // 60))  # rough page estimate
    return text, pages, warnings


def supports_mime(mime_type: str) -> bool:
    """Check whether this processor can handle the given MIME type."""
    if mime_type == "application/pdf":
        return True
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return True
    return bool(mime_type.startswith("text/"))
